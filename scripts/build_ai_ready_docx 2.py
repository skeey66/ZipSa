# -*- coding: utf-8 -*-
"""§2 AI-Ready (프롬프트/JSON 스키마) → docs/ZipSa-AI-Ready.docx

    python3 scripts/build_ai_ready_docx.py

제출 가이드(제출 가이드.pdf) §4-1 설계문서 §2 AI-Ready 절 하나만 뽑아 Word 문서로 만듭니다.
내용은 실제 코드(ai/LlmInsight.java, docs/api/API.yml)에서 그대로 옮겼습니다 — 지어낸 문장이 없습니다.
"""
import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
OUT = ROOT / "docs/ZipSa-AI-Ready.docx"

INK = RGBColor(0x11, 0x19, 0x27)
BODY = RGBColor(0x33, 0x41, 0x5A)
MUTE = RGBColor(0x7C, 0x89, 0x9B)
ACCENT = RGBColor(0x1B, 0x24, 0x30)
CODE_BG = "F3F5F9"
HEAD_BG = "1B2430"
MONO = "Consolas"
SANS = "맑은 고딕"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="D7DEE8", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), SANS)

    for sec in doc.sections:
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ZIP 보금자리 · ZipSa — 설계문서 §2")
    r.font.size = Pt(10)
    r.font.color.rgb = MUTE
    r.font.bold = True

    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run("AI-Ready — 프롬프트(시스템/유저) 실제 문장, 입출력 JSON 스키마 예시")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = INK

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    r = sub.add_run(
        "실제 코드 원문 그대로입니다 — 시스템/유저 프롬프트는 backend/src/main/java/com/zipsa/ai/LlmInsight.java, "
        "JSON 스키마 필드명은 docs/api/API.yml · docs/DB.dbml 과 동일합니다."
    )
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = MUTE

    # 구분선
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:color"), "1B2430")
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(14)


def add_h2(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "1B2430")
    left.set(qn("w:space"), "8")
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(f"{num}  {text}")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = INK


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = ACCENT


def add_body(doc, text, size=10, color=None, space_after=8, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color or BODY
    r.font.italic = italic
    return p


def add_code_block(doc, text, size=9):
    """모노스페이스 + 옅은 배경 박스로 프롬프트/JSON 원문을 그대로 보여준다."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    set_cell_borders(cell)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("top", "120"), ("bottom", "120"), ("left", "160"), ("right", "160")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

    cell.paragraphs[0].text = ""
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(line if line else " ")
        r.font.name = MONO
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0x1F, 0x33, 0x50)
        rpr = r._element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), MONO)
        rFonts.set(qn("w:ascii"), MONO)
        rFonts.set(qn("w:hAnsi"), MONO)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)
    spacer.paragraph_format.space_before = Pt(0)
    spacer_run = spacer.add_run("")
    spacer_run.font.size = Pt(2)


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x5B, 0x6B, 0x80)
    r.font.small_caps = True


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], HEAD_BG)
        set_cell_borders(hdr[i], color="1B2430")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = BODY
            set_cell_borders(cells[i])
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)
    sp.add_run("").font.size = Pt(2)
    return table


# ── 문서 ─────────────────────────────────────────────────
doc = Document()
style_base(doc)
add_title_block(doc)

# ── ① AI 확장 지점 ─────────────────────────────────────────
add_h2(doc, "①", "AI 확장 지점")
add_body(
    doc,
    "정책 상세 화면(SCR-POLICY-002)에서 “AI 요약 보기”를 누르면 "
    "GET /api/ai/policies/{policyId} 를 호출한다. 이 호출이 AI가 실제로 개입하는 지점이며, "
    "응답으로 정책 요약 3줄과 “나에게 적용되는지” 카드(결론·근거·다음 행동)를 렌더링한다.",
)

# ── ② 프롬프트 실제 문장 ────────────────────────────────────
add_h2(doc, "②", "프롬프트 실제 문장 (시스템 / 유저)")

add_h3(doc, "시스템 프롬프트")
add_code_block(doc, """너는 한국의 청년 주거 서비스 'ZIP 보금자리'의 정책 안내 도우미다.
사용자에게 존댓말로, 군더더기 없이 사실만 전달한다.

반드시 지킬 것
- 제공된 사실만 사용한다. 없는 금액·날짜·조건을 지어내지 않는다.
- 이미 계산된 판정(신청 가능 여부, 남은 일수)을 그대로 따른다. 다시 계산하지 않는다.
- summary 는 3문장. 각 문장은 한 줄로 끝낸다.
- verdict(한 줄 결론)는 이미 확정돼 있다. 바꾸거나 뒤집지 않는다.
- reasons 는 판정의 근거를 2~4개.
- nextSteps 는 사용자가 지금 할 수 있는 구체적인 행동을 1~3개.
- 과장하거나 권유하지 않는다. "꼭 신청하세요" 같은 표현은 쓰지 않는다.""")

add_h3(doc, "유저 프롬프트 (예시 값을 채운 실제 전송본)")
add_code_block(doc, """아래 청년 정책을 요약하고, 이 회원에게 어떻게 적용되는지 설명해 줘.

[정책]
제목: 청년 월세 특별지원
분류: 주거지원
대상 지역: 서울
주관: 서울특별시
지원 나이: 만 19~34세
소득 조건: 연 5천만원 이하
신청 기간: 2026-08-01 ~ 2026-09-20 (마감까지 17일)
신청 방법: 복지로 온라인 신청
내용:
서울시에 거주하는 무주택 청년 1인가구에게 월세를 최대 20만원씩 12개월간 지원합니다.
기준중위소득 150% 이하이며 임차보증금 5천만원, 월세 60만원 이하 주택에 거주해야 합니다.

[회원]
20대 후반 · 직장인 · 연소득 3~4천만원 · 미혼 · 거주 서울 관악구

[이미 확정된 판정 — 이대로 따를 것]
결론: 신청 조건에 해당합니다. 마감 전에 서류를 준비하세요.
판정 근거: 나이대(20대 후반)가 대상 범위에 듭니다.
- 연소득 3,000~4,000만원이 4천만원 이하 조건을 만족합니다.
- 거주지가 관악구로 일치합니다.""")

# ── ③ 입출력 JSON 스키마 예시 ───────────────────────────────
add_h2(doc, "③", "입출력 JSON 스키마 예시")
add_body(
    doc,
    "필드명은 API.yml 의 AiInsightResponse 스키마, DB.dbml 의 policies 테이블과 동일하게 맞췄다 "
    "— 프론트·백엔드·AI 세 곳이 같은 이름을 쓰므로 매핑 오류 없이 연동된다.",
    italic=True,
)

add_label(doc, "입력 (요청) — GET /api/ai/policies/{policyId}")
add_code_block(doc, """{
  "policyId": 1024,
  "policy": {
    "title": "청년 월세 특별지원",
    "category": "HOUSING",
    "region": "서울",
    "targetAgeRange": "만 19~34세",
    "targetSalaryRange": "연 5천만원 이하",
    "applyEndDate": "2026-09-20",
    "applyMethod": "복지로 온라인 신청"
  },
  "user": {
    "ageRange": "AGE_20S_LATE",
    "job": "EMPLOYEE",
    "salaryRange": "RANGE_3000_4000",
    "maritalStatus": "SINGLE",
    "region": "서울 관악구"
  },
  "verdict": {
    "headline": "신청 조건에 해당합니다. 마감 전에 서류를 준비하세요.",
    "facts": [
      "나이대(20대 후반)가 대상 범위에 듭니다",
      "연소득 3,000~4,000만원이 4천만원 이하 조건을 만족합니다",
      "거주지가 관악구로 일치합니다"
    ],
    "tone": "good"
  }
}""", size=8.6)

add_label(doc, "출력 (응답) — AiInsightResponse")
add_code_block(doc, """{
  "success": true,
  "data": {
    "summary": [
      "청년 월세 특별지원은 만 19~34세 무주택 청년에게 월세를 최대 20만원씩 12개월 지원하는 사업입니다.",
      "서울시 거주자 대상이며 기준중위소득 150% 이하, 임차보증금 5천만원 이하 주택이 조건입니다.",
      "신청은 복지로 홈페이지에서 온라인으로 진행하며 마감일까지 서류 제출을 완료해야 합니다."
    ],
    "application": {
      "verdict": "신청 조건에 해당합니다. 마감 전에 서류를 준비하세요.",
      "reasons": [
        "나이대(20대 후반)가 대상 범위에 듭니다",
        "연소득 3,000~4,000만원이 4천만원 이하 조건을 만족합니다",
        "거주지가 관악구로 일치합니다"
      ],
      "nextSteps": [
        "무주택 확인을 위한 주민등록등본 발급",
        "임대차계약서 사본 준비",
        "정부24에서 온라인 신청"
      ],
      "tone": "good"
    },
    "aiGenerated": true
  },
  "error": null
}""", size=8.6)

add_h3(doc, "필드명 매핑 — API.yml · DB.dbml 과 이름 일치")
add_table(
    doc,
    ["JSON 필드", "API.yml (AiInsightResponse)", "DB.dbml (policies)"],
    [
        ["summary", "summary: string[]", "— (LLM 생성, 저장 안 함)"],
        ["application.verdict", "verdict: string", "— (RuleBasedInsight 계산값)"],
        ["application.reasons", "reasons: string[]", "policies.target_age_range 등에서 파생"],
        ["application.nextSteps", "nextSteps: string[]", "policies.apply_method 에서 파생"],
        ["aiGenerated", "aiGenerated: boolean", "—"],
        ["policyId", "path parameter {policyId}", "policies.id (PK)"],
    ],
    col_widths=[4.2, 5.6, 6.0],
)

add_body(
    doc,
    "(선택·심화) 응답 지연에 대비해 aiGenerated 옆에 status 필드(pending / completed / failed)를 "
    "추가하면 가점 대상이다. 현재 구현은 동기 호출이라 이 필드가 없다.",
    size=9.5,
    color=MUTE,
    italic=True,
)

doc.save(OUT)
print(f"{OUT.relative_to(ROOT)} — {OUT.stat().st_size // 1024}KB")
